import os
import io
import json
import uuid
import shutil
from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # 500MB max
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')
app.config['OUTPUT_FOLDER'] = os.path.join(os.path.dirname(__file__), 'output')

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload_images():
    """Handle multiple image uploads."""
    if 'images' not in request.files:
        return jsonify({'error': 'No images provided'}), 400

    files = request.files.getlist('images')
    uploaded = []

    for f in files:
        if f and f.filename:
            ext = os.path.splitext(f.filename)[1].lower()
            if ext not in ('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.webp'):
                continue
            unique_name = f"{uuid.uuid4().hex}{ext}"
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_name)
            f.save(filepath)

            # Create thumbnail for preview
            try:
                img = Image.open(filepath)
                img.thumbnail((400, 400), Image.LANCZOS)
                thumb_name = f"thumb_{unique_name}"
                # Save as JPEG for smaller size
                thumb_path = os.path.join(app.config['UPLOAD_FOLDER'], thumb_name)
                if img.mode in ('RGBA', 'P'):
                    img = img.convert('RGB')
                img.save(thumb_path, 'JPEG', quality=80)
            except Exception:
                thumb_name = unique_name

            uploaded.append({
                'filename': unique_name,
                'thumb': thumb_name,
                'original_name': f.filename
            })

    return jsonify({'uploaded': uploaded})


@app.route('/uploads/<filename>')
def serve_upload(filename):
    """Serve uploaded files."""
    return send_file(os.path.join(app.config['UPLOAD_FOLDER'], filename))


@app.route('/generate', methods=['POST'])
def generate_docx():
    """Generate a Word document from specimen data."""
    data = request.get_json()
    specimens = data.get('specimens', [])
    title = data.get('title', 'Weed Specimen Record')
    cols_per_row = data.get('cols_per_row', 3)

    if not specimens:
        return jsonify({'error': 'No specimens provided'}), 400

    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # ── Title ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    # ── Build table in groups of cols_per_row ──
    usable_width = section.page_width - section.left_margin - section.right_margin
    col_width = usable_width // cols_per_row

    for group_start in range(0, len(specimens), cols_per_row):
        group = specimens[group_start:group_start + cols_per_row]
        actual_cols = len(group)

        # Create table: 3 rows (header, image, details) x actual_cols columns
        table = doc.add_table(rows=3, cols=actual_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Set column widths
        for col_idx in range(actual_cols):
            for row in table.rows:
                row.cells[col_idx].width = col_width

        # Style the table borders
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
            '  <w:left w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
            '  <w:right w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
            '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
            '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
            '</w:tblBorders>'
        )
        tblPr.append(borders)

        for col_idx, spec in enumerate(group):
            specimen_num = group_start + col_idx + 1
            label = spec.get('label', f'Specimen {specimen_num}')

            # ── Row 0: Header (green background) ──
            header_cell = table.rows[0].cells[col_idx]
            header_cell.text = ''
            p = header_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(label)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Green background
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="2E7D32" w:val="clear"/>'
            )
            header_cell._tc.get_or_add_tcPr().append(shading)

            # Cell padding
            _set_cell_padding(header_cell, top=60, bottom=60)

            # ── Row 1: Image ──
            img_cell = table.rows[1].cells[col_idx]
            img_cell.text = ''
            p = img_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            img_filename = spec.get('filename', '')
            img_path = os.path.join(app.config['UPLOAD_FOLDER'], img_filename)
            if os.path.exists(img_path):
                try:
                    # Resize image for Word
                    pil_img = Image.open(img_path)
                    if pil_img.mode in ('RGBA', 'P'):
                        pil_img = pil_img.convert('RGB')
                    # Fit within column
                    max_w = 500
                    max_h = 400
                    pil_img.thumbnail((max_w, max_h), Image.LANCZOS)
                    buf = io.BytesIO()
                    pil_img.save(buf, format='JPEG', quality=85)
                    buf.seek(0)

                    target_width = Cm(5.5) if cols_per_row == 3 else Cm(7)
                    run = p.add_run()
                    run.add_picture(buf, width=target_width)
                except Exception as e:
                    run = p.add_run(f'[Image Error: {str(e)}]')
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            else:
                run = p.add_run('[No Image]')
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            _set_cell_padding(img_cell, top=80, bottom=80)

            # ── Row 2: Details ──
            detail_cell = table.rows[2].cells[col_idx]
            detail_cell.text = ''

            details = [
                ('Common Name', spec.get('common_name', '')),
                ('Scientific Name', spec.get('scientific_name', '')),
                ('Family', spec.get('family', '')),
                ('Type', spec.get('type', '')),
            ]
            notes = spec.get('notes', '')
            if notes:
                details.append(('Notes', notes))

            for i, (key, val) in enumerate(details):
                if i == 0:
                    p = detail_cell.paragraphs[0]
                else:
                    p = detail_cell.add_paragraph()
                p.space_before = Pt(1)
                p.space_after = Pt(1)

                run_key = p.add_run(f'{key} – ')
                run_key.bold = True
                run_key.font.size = Pt(9)
                run_key.font.name = 'Arial'
                run_key.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

                run_val = p.add_run(val if val else '—')
                run_val.font.size = Pt(9)
                run_val.font.name = 'Arial'
                if key == 'Scientific Name' and val:
                    run_val.italic = True
                run_val.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            _set_cell_padding(detail_cell, top=60, bottom=60, left=80, right=80)

        # Add spacing between groups
        spacer = doc.add_paragraph()
        spacer.space_before = Pt(6)
        spacer.space_after = Pt(6)

    # Save document
    output_name = f"specimen_record_{uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_name)
    doc.save(output_path)

    return jsonify({'filename': output_name})


@app.route('/download/<filename>')
def download_file(filename):
    """Download generated Word file."""
    filepath = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404
    return send_file(
        filepath,
        as_attachment=True,
        download_name='Weed_Specimen_Record.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/clear', methods=['POST'])
def clear_all():
    """Clear all uploaded files and generated documents."""
    for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
        for f in os.listdir(folder):
            fp = os.path.join(folder, f)
            if os.path.isfile(fp):
                os.remove(fp)
    return jsonify({'status': 'cleared'})


@app.route('/save-project', methods=['POST'])
def save_project():
    """Save project data as JSON."""
    data = request.get_json()
    save_path = os.path.join(app.config['OUTPUT_FOLDER'], 'project_save.json')
    with open(save_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2)
    return jsonify({'status': 'saved'})


@app.route('/load-project', methods=['GET'])
def load_project():
    """Load saved project data."""
    save_path = os.path.join(app.config['OUTPUT_FOLDER'], 'project_save.json')
    if not os.path.exists(save_path):
        return jsonify({'error': 'No saved project found'}), 404
    with open(save_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return jsonify(data)


def _set_cell_padding(cell, top=0, bottom=0, left=40, right=40):
    """Set cell padding in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


if __name__ == '__main__':
    app.run(debug=True, port=5000)
